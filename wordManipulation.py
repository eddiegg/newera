import docx,logging,subprocess
logging.basicConfig(level=logging.DEBUG,format='%(levelname)s - %(message)s')
doc = docx.Document('./test/demo.docx')
logging.debug(len(doc.paragraphs))
logging.debug(len(doc.paragraphs))
logging.debug(doc.paragraphs[2].runs[1].text)

subprocess.Popen(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe")
